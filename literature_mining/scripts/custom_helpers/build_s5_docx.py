"""
build_s5_docx.py — Assemble Supplementary Section S5 (.docx).

Pulls every Phase 2/3 artifact (counts matrix, ratio, top-cited tables, topic
analysis, figures, abstracts) into one supplementary document following the plan's
S5.1–S5.7 structure.

Output: reports/supplementary_drafts/HIF1A_COVID_Supplementary_S5.docx
        reports/supplementary_drafts/TableS5_1_counts.csv
        reports/supplementary_drafts/TableS5_4_hif1a_top.csv
        reports/supplementary_drafts/TableS5_5_aivaccine_top.csv
"""

import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "modified_framework"))
import covid_config as cfg

SUPP = cfg.REPORTS_DIR / "supplementary_drafts"
SUPP.mkdir(parents=True, exist_ok=True)
TAG_RE = re.compile(r"<[^>]+>")

# abstracts library sizes — full 140 (matches PRISMA: G3/G5 top 50, G2/G4 top 20)
ABS_PER_GROUP = {"G3": 50, "G5": 50, "G2": 20, "G4": 20}
THEME_LABELS = {
    1: "Healthcare systems, workforce & societal impact",
    2: "Long COVID & clinical outcomes",
    3: "Molecular virology & host immunity",
    4: "Mental health",
    5: "Vaccines & mRNA technology",
    6: "Surveillance & co-circulating respiratory viruses",
}


def clean(text):
    return re.sub(r"\s+", " ", TAG_RE.sub("", str(text))).strip()


def dedup_top(df):
    df = df.copy()
    df["tn"] = df["title"].str.lower().str.replace(r"[^a-z0-9 ]", "", regex=True).str.strip()
    return df.drop_duplicates("tn", keep="first").drop(columns="tn")


# ----------------------------- docx helpers -----------------------------
def h(doc, text, level):
    doc.add_heading(text, level=level)


def para(doc, text, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    return p


def add_figure(doc, path, cap, width=6.2):
    if not Path(path).exists():
        para(doc, f"[missing figure: {path}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def add_table(doc, df, cols=None, headers=None, widths=None, fontsize=8):
    cols = cols or list(df.columns)
    headers = headers or cols
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, hd in enumerate(headers):
        hdr[i].text = str(hd)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(fontsize)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = clean(row[c])
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(fontsize)
    return t


# ----------------------------- build tables -----------------------------
def table_s5_1():
    m = pd.read_csv(cfg.RAW_DIR / "covid_lit_matrix.csv")
    r = pd.read_csv(cfg.RAW_DIR / "covid_ratio_by_year.csv")[["Year", "ai_ratio_pct"]]
    df = m.merge(r, on="Year")
    df.columns = ["Year", "G1 COVID", "G2 AI", "G3 HIF1A", "G5 AI-vax", "AI ratio %"]
    for c in ["G1 COVID", "G2 AI", "G3 HIF1A", "G5 AI-vax"]:
        df[c] = df[c].map(lambda x: f"{int(x):,}")
    df["AI ratio %"] = df["AI ratio %"].map(lambda x: f"{x:.2f}")
    df.to_csv(SUPP / "TableS5_1_counts.csv", index=False)
    return df


def table_topcited(group, n):
    # read the full fetched set so dedup yields a true top-N unique (the _top.csv
    # was pre-truncated to 50, so dedup there would drop below N)
    df = pd.read_csv(cfg.CITED_DIR / f"{group}_all.csv")
    df = df.sort_values("citations", ascending=False)
    df = dedup_top(df).head(n).copy()
    df["rank"] = range(1, len(df) + 1)
    df["title"] = df["title"].map(clean)
    df["citations"] = df["citations"].map(lambda x: f"{int(x):,}")
    out = df[["rank", "citations", "year", "journal", "title", "doi"]]
    return out


# ----------------------------- main build -----------------------------
def main():
    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = doc.add_heading("Supplementary Section S5 — COVID-19 Literature Mining Analysis", level=0)

    para(doc,
         "This section reports an automated bibliometric analysis of the COVID-19 "
         "literature performed to contextualize the present multi-scale HIF1A study. "
         "Five PubMed query groups (G1–G5) were retrieved over 2019/01/01–2026/06/01 "
         "using the octe-ai-mining-framework and bioprinting-ai-mining-framework, with "
         "citation ranking via OpenAlex. All counts are deduplicated single-range queries "
         "unless stated otherwise.")

    qs = cfg.queries(decorated=True)
    val = pd.read_csv(cfg.RAW_DIR / "query_validation_counts.csv").set_index("group")

    # ---- S5.1 ----
    h(doc, "S5.1  Search methodology and strategy", 1)
    para(doc,
         "Literature retrieval used NCBI Entrez (Biopython) with an institutional API key. "
         "For each group, publication counts were obtained per calendar year and as a "
         "single deduplicated date-range query; per-year sums slightly exceed the range "
         "query because PubMed indexes some records under both electronic and print "
         "publication years. Top-cited papers were ranked by OpenAlex cited_by_count "
         "(matched by DOI). The COVID-19 base was defined as the MeSH terms \"COVID-19\" or "
         "\"SARS-CoV-2\" or their occurrence in title/abstract; no publication-type or "
         "language restriction was applied. Date range: 2019/01/01–2026/06/01.")
    qrows = []
    for g in ["G1", "G2", "G3", "G4", "G5"]:
        qrows.append({"Group": g, "Label": cfg.GROUP_LABELS[g],
                      "N": f"{int(val.loc[g,'n']):,}", "PubMed query": qs[g]})
    add_table(doc, pd.DataFrame(qrows), cols=["Group", "Label", "N", "PubMed query"],
              fontsize=7)
    caption(doc, "Table S5.0 — PubMed query strings and total records per group "
                 "(bounded 2019/01/01–2026/06/01, deduplicated). G4 uses the 2024–2026 window.")

    # ---- S5.2 PRISMA ----
    h(doc, "S5.2  PRISMA-like flow diagram", 1)
    add_figure(doc, cfg.FIGURES_DIR / "prisma" / "FigS5_5_prisma.png",
               "Figure S5.5 — PRISMA-like flow of the literature mining process.")

    # ---- S5.3 Trend analysis ----
    h(doc, "S5.3  Publication trends and the AI ratio", 1)
    t1 = table_s5_1()
    add_table(doc, t1, fontsize=8)
    caption(doc, "Table S5.1 — COVID-19 publication counts by year and group, with the AI "
                 "ratio (G2/G1). Per-year counts; 2026 is partial (through 2026/06/01).")
    add_figure(doc, cfg.FIGURES_DIR / "trends" / "FigS5_1_growth.png",
               "Figure S5.1 — Annual COVID-19 publications (a) and AI/AI-vaccine/HIF1A "
               "subsets on a log scale (b). 2026 partial.")
    add_figure(doc, cfg.FIGURES_DIR / "heatmaps" / "FigS5_2_heatmap.png",
               "Figure S5.2 — Publication intensity by group and year (row-normalized; "
               "cells show raw counts).")
    add_figure(doc, cfg.FIGURES_DIR / "ratios" / "FigS5_3_ai_ratio.png",
               "Figure S5.3 — AI/ML publications as a percentage of COVID-19 publications "
               "per year (overall 2.70%); the share accelerates after the 2022 LLM era.")

    # ---- S5.4 Recent topics ----
    h(doc, "S5.4  Recent topics (2024–2026)", 1)
    para(doc, "The 2024–2026 corpus (2,896 abstracts of the most-cited recent papers) was "
              "analyzed by TF-IDF term weighting and unsupervised K-means clustering "
              "(k=6). Domain-ubiquitous terms were removed before vectorization.")
    tt = pd.read_csv(cfg.TOPIC_DIR / "top_terms.csv").head(30)
    tt2 = tt[["rank", "term", "doc_pct"]].copy()
    tt2["doc_pct"] = tt2["doc_pct"].map(lambda x: f"{x:.1f}%")
    add_table(doc, tt2, headers=["Rank", "Term", "% of abstracts"], fontsize=8)
    caption(doc, "Table S5.2 — Top 30 distinctive terms in 2024–2026 abstracts "
                 "(ranked by aggregate TF-IDF weight).")
    add_figure(doc, cfg.FIGURES_DIR / "wordclouds" / "FigS5_4_wordcloud.png",
               "Figure S5.4 — Word cloud of distinctive 2024–2026 terms "
               "(domain-generic words removed).")
    tc = pd.read_csv(cfg.TOPIC_DIR / "topic_clusters.csv")
    tc["theme"] = tc["rank"].map(THEME_LABELS)
    tc["pct"] = tc["pct"].map(lambda x: f"{x:.1f}%")
    add_table(doc, tc, cols=["rank", "theme", "n_docs", "pct", "top_terms"],
              headers=["#", "Theme", "n", "Share", "Characteristic terms"], fontsize=7.5)
    caption(doc, "Table S5.3 — Topic clusters in the 2024–2026 literature. Cluster 1 mixes "
                 "in weakly COVID-specific literature; clusters 2–6 carry the substantive signal.")

    # ---- S5.5 HIF1A ----
    h(doc, "S5.5  HIF1A focus and gap analysis", 1)
    para(doc, "Only 139 HIF1A-related COVID-19 publications were indexed across 2019–2026 "
              "(47 in 2024–2026) — a small, slowly growing niche. The most-cited papers are "
              "predominantly single-scale mechanistic studies establishing the HIF-1α↔"
              "SARS-CoV-2 axis (glycolytic reprogramming, ORF3a/ROS-driven HIF-1α "
              "stabilization, ACE2 regulation, and the hypoxia/cytokine-storm link). None of "
              "the retrieved studies integrate populational, transcriptomic, single-cell, "
              "clinical, and host-genetic evidence within a unified predictive framework — "
              "the gap addressed by the present multi-scale analysis.")
    t4 = table_topcited("G3", 50)
    t4.to_csv(SUPP / "TableS5_4_hif1a_top.csv", index=False)
    add_table(doc, t4, headers=["#", "Cites", "Year", "Journal", "Title", "DOI"], fontsize=7)
    caption(doc, "Table S5.4 — Top 50 HIF1A–COVID-19 publications by OpenAlex citation count.")

    # ---- S5.6 AI-vaccine ----
    h(doc, "S5.6  AI-assisted vaccine design outlook", 1)
    para(doc, "Group G5 (COVID-19 + AI + vaccine; 1,090 papers) captures a heterogeneous "
              "mixture of general AI applications and genuine AI-driven vaccine-design work. "
              "Machine-learning reverse vaccinology (Vaxign-ML) and in-silico deep-learning "
              "multi-epitope design are the most directly relevant COVID-specific examples, "
              "complementing general-purpose tools (AlphaFold, ProteinMPNN, LinearDesign) "
              "discussed in the Conclusions.")
    t5 = table_topcited("G5", 30)
    t5.to_csv(SUPP / "TableS5_5_aivaccine_top.csv", index=False)
    add_table(doc, t5, headers=["#", "Cites", "Year", "Journal", "Title", "DOI"], fontsize=7)
    caption(doc, "Table S5.5 — Top 30 COVID-19 AI/vaccine publications by citation count.")

    # ---- S5.7 Abstracts library ----
    h(doc, "S5.7  Selected abstracts library", 1)
    para(doc, "Curated abstracts of the most-cited papers per group (HIF1A and AI-vaccine "
              "prioritized). Full top-cited tables appear above; complete abstracts for all "
              "140 retrieved papers and a BibTeX file accompany this supplement.")
    for g, n in ABS_PER_GROUP.items():
        h(doc, f"S5.7.{g} — {cfg.GROUP_LABELS[g]} (top {n})", 2)
        corp = pd.read_csv(cfg.CITED_DIR / f"{g}_corpus.csv")
        corp = dedup_top(corp.sort_values("citations", ascending=False)).head(n)
        for i, (_, row) in enumerate(corp.iterrows(), 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {clean(row['title'])}")
            r.bold = True
            r.font.size = Pt(9.5)
            meta = doc.add_paragraph()
            mr = meta.add_run(f"{row['journal']} ({row['year']}) · "
                              f"{int(row['citations']):,} citations · doi:{row['doi']}")
            mr.italic = True
            mr.font.size = Pt(8.5)
            mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            ab = clean(row["abstract"])
            abp = doc.add_paragraph(ab if ab else "[no abstract available]")
            for rr in abp.runs:
                rr.font.size = Pt(9)

    out = SUPP / "HIF1A_COVID_Supplementary_S5.docx"
    doc.save(str(out))
    print(f"Saved: {out}")
    print(f"  paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")
    # also report file size
    print(f"  size: {out.stat().st_size/1024:.0f} KB")


if __name__ == "__main__":
    main()
